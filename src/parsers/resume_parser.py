"""
Resume Parser Module
Extracts information from PDF and DOCX resume files
"""

import logging
import re
import fitz  # PyMuPDF
from pathlib import Path
from typing import Dict, List, Optional, Union
from docx import Document
import spacy
from dataclasses import dataclass

from config import config

logger = logging.getLogger(__name__)

@dataclass
class ResumeData:
    """Data class to hold parsed resume information"""
    raw_text: str
    name: str
    email: str
    phone: str
    summary: str
    skills: List[str]
    experience: List[Dict[str, str]]
    education: List[Dict[str, str]]
    sections: Dict[str, str]

class ResumeParser:
    """Parser for extracting information from resume files"""
    
    def __init__(self):
        """Initialize the resume parser"""
        self.supported_formats = ['.pdf', '.docx', '.doc']
        
        # Load spaCy model for NLP processing
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            logger.warning("spaCy model 'en_core_web_sm' not found. Install with: python -m spacy download en_core_web_sm")
            self.nlp = None
        
        # Common section headers
        self.section_headers = {
            'summary': ['summary', 'profile', 'objective', 'about', 'overview'],
            'experience': ['experience', 'work experience', 'employment', 'work history', 'professional experience'],
            'education': ['education', 'academic background', 'qualifications'],
            'skills': ['skills', 'technical skills', 'core competencies', 'expertise', 'technologies'],
            'projects': ['projects', 'key projects', 'notable projects'],
            'certifications': ['certifications', 'certificates', 'licenses']
        }
        
        # Common skill keywords
        self.skill_keywords = [
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust',
            'swift', 'kotlin', 'scala', 'r', 'matlab', 'sql', 'html', 'css',
            
            # Frameworks & Libraries
            'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'spring', 'laravel',
            'tensorflow', 'pytorch', 'pandas', 'numpy', 'scikit-learn', 'opencv',
            
            # Databases
            'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'oracle', 'sqlite',
            
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'github', 'gitlab',
            'terraform', 'ansible', 'linux', 'unix',
            
            # Other Technologies
            'api', 'rest', 'graphql', 'microservices', 'agile', 'scrum', 'machine learning',
            'data science', 'artificial intelligence', 'blockchain'
        ]
    
    def parse_resume(self, file_path: Union[str, Path]) -> ResumeData:
        """
        Parse resume file and extract information
        
        Args:
            file_path: Path to resume file
            
        Returns:
            ResumeData object with extracted information
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")
        
        if file_path.suffix.lower() not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        logger.info(f"Parsing resume: {file_path}")
        
        # Extract text based on file type
        if file_path.suffix.lower() == '.pdf':
            raw_text = self._extract_text_from_pdf(file_path)
        else:  # .docx or .doc
            raw_text = self._extract_text_from_docx(file_path)
        
        # Parse extracted text
        return self._parse_text(raw_text)
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            doc = fitz.open(file_path)
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text += page.get_text()
            
            doc.close()
            return text
        
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise
    
    def _parse_text(self, text: str) -> ResumeData:
        """Parse extracted text and extract structured information"""
        
        # Clean text
        cleaned_text = self._clean_text(text)
        
        # Extract basic information
        name = self._extract_name(cleaned_text)
        email = self._extract_email(cleaned_text)
        phone = self._extract_phone(cleaned_text)
        
        # Extract sections
        sections = self._extract_sections(cleaned_text)
        
        # Extract specific information
        summary = self._extract_summary(sections)
        skills = self._extract_skills(cleaned_text, sections)
        experience = self._extract_experience(sections)
        education = self._extract_education(sections)
        
        return ResumeData(
            raw_text=text,
            name=name,
            email=email,
            phone=phone,
            summary=summary,
            skills=skills,
            experience=experience,
            education=education,
            sections=sections
        )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s@.-]', ' ', text)
        return text.strip()
    
    def _extract_name(self, text: str) -> str:
        """Extract candidate name from resume"""
        lines = text.split('\n')
        
        # Usually name is in the first few lines
        for line in lines[:5]:
            line = line.strip()
            if len(line) > 2 and len(line.split()) <= 4:
                # Check if it looks like a name (no numbers, not email)
                if not re.search(r'\d|@', line) and len(line) > 5:
                    return line
        
        return "Unknown"
    
    def _extract_email(self, text: str) -> str:
        """Extract email address from resume"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        return emails[0] if emails else ""
    
    def _extract_phone(self, text: str) -> str:
        """Extract phone number from resume"""
        phone_patterns = [
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',  # 123-456-7890 or 123.456.7890
            r'\(\d{3}\)\s*\d{3}[-.]?\d{4}',    # (123) 456-7890
            r'\b\d{10}\b'                       # 1234567890
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                return phones[0]
        
        return ""
    
    def _extract_sections(self, text: str) -> Dict[str, str]:
        """Extract different sections from resume"""
        sections = {}
        lines = text.split('\n')
        current_section = 'header'
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            section_found = None
            for section_type, headers in self.section_headers.items():
                for header in headers:
                    if header.lower() in line.lower() and len(line) < 50:
                        section_found = section_type
                        break
                if section_found:
                    break
            
            if section_found:
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                
                # Start new section
                current_section = section_found
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def _extract_summary(self, sections: Dict[str, str]) -> str:
        """Extract professional summary"""
        for section_name in ['summary', 'profile', 'objective', 'about']:
            if section_name in sections:
                return sections[section_name]
        
        # If no dedicated summary section, use first few lines
        if 'header' in sections:
            lines = sections['header'].split('\n')
            summary_lines = []
            for line in lines:
                if len(line) > 50 and not re.search(r'@|\d{3}[-.]?\d{3}', line):
                    summary_lines.append(line)
                    if len(summary_lines) >= 3:
                        break
            return '\n'.join(summary_lines)
        
        return ""
    
    def _extract_skills(self, text: str, sections: Dict[str, str]) -> List[str]:
        """Extract skills from resume"""
        skills = set()
        
        # Extract from skills section
        if 'skills' in sections:
            skills_text = sections['skills'].lower()
            for skill in self.skill_keywords:
                if skill in skills_text:
                    skills.add(skill.title())
        
        # Extract from entire text
        text_lower = text.lower()
        for skill in self.skill_keywords:
            if skill in text_lower:
                skills.add(skill.title())
        
        return sorted(list(skills))
    
    def _extract_experience(self, sections: Dict[str, str]) -> List[Dict[str, str]]:
        """Extract work experience"""
        experience = []
        
        if 'experience' in sections:
            exp_text = sections['experience']
            # Simple extraction - can be enhanced with more sophisticated parsing
            lines = exp_text.split('\n')
            current_job = {}
            
            for line in lines:
                line = line.strip()
                if not line:
                    if current_job:
                        experience.append(current_job)
                        current_job = {}
                    continue
                
                # Try to identify job titles, companies, dates
                if not current_job.get('title'):
                    current_job['title'] = line
                elif not current_job.get('company'):
                    current_job['company'] = line
                else:
                    if 'description' not in current_job:
                        current_job['description'] = line
                    else:
                        current_job['description'] += ' ' + line
            
            if current_job:
                experience.append(current_job)
        
        return experience
    
    def _extract_education(self, sections: Dict[str, str]) -> List[Dict[str, str]]:
        """Extract education information"""
        education = []
        
        if 'education' in sections:
            edu_text = sections['education']
            lines = edu_text.split('\n')
            
            for line in lines:
                line = line.strip()
                if line and len(line) > 10:
                    education.append({
                        'degree': line,
                        'institution': '',
                        'year': ''
                    })
        
        return education
