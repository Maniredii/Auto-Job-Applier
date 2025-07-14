# Author: Enhanced by AI Assistant
# Resume Optimization Engine
# AI-powered resume customization system that tailors resumes for each job application

import os
import json
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from modules.helpers import print_lg

# Import Groq AI for resume optimization
try:
    from modules.ai.groqConnections import groq_create_client
    from config.secrets import groq_api_key, groq_model
    GROQ_AVAILABLE = True
except ImportError:
    GROQ_AVAILABLE = False
    print_lg("⚠️ Groq AI not available for resume optimization")

@dataclass
class ResumeSection:
    """Data class for resume sections."""
    title: str
    content: List[str]
    priority: int = 1
    keywords: List[str] = None

@dataclass
class JobRequirements:
    """Data class for job requirements analysis."""
    required_skills: List[str]
    preferred_skills: List[str]
    experience_level: str
    education_requirements: List[str]
    keywords: List[str]
    industry: str
    job_function: str

class ResumeOptimizer:
    """
    AI-powered resume optimization engine that customizes resumes for specific job applications.
    """
    
    def __init__(self, base_resume_path: str = "all resumes/default/resume.docx"):
        self.base_resume_path = base_resume_path
        self.optimized_resumes_dir = "all resumes/optimized"
        self.templates_dir = "templates/resume_templates"

        # Ensure directories exist
        os.makedirs(self.optimized_resumes_dir, exist_ok=True)
        os.makedirs(self.templates_dir, exist_ok=True)

        # Load skill mappings and synonyms
        self.skill_mappings = self._load_skill_mappings()
        self.keyword_weights = self._load_keyword_weights()

        # Resume parsing cache
        self.parsed_resume_cache = None

        # Initialize Groq AI client for resume optimization
        self.groq_client = None
        if GROQ_AVAILABLE and groq_api_key:
            try:
                self.groq_client = groq_create_client(groq_api_key, groq_model)
                if self.groq_client:
                    print_lg("🚀 Groq AI enabled for resume optimization")
            except Exception as e:
                print_lg(f"⚠️ Failed to initialize Groq AI: {e}")
                self.groq_client = None
        
    def _load_skill_mappings(self) -> Dict:
        """Load skill mappings and synonyms."""
        default_mappings = {
            "programming_languages": {
                "python": ["python", "py", "python3", "django", "flask", "fastapi"],
                "javascript": ["javascript", "js", "node.js", "nodejs", "react", "vue", "angular"],
                "java": ["java", "spring", "spring boot", "hibernate"],
                "sql": ["sql", "mysql", "postgresql", "sqlite", "database"],
                "r": ["r programming", "r language", "rstudio"],
                "c++": ["c++", "cpp", "c plus plus"],
                "c#": ["c#", "csharp", "c sharp", ".net", "dotnet"]
            },
            "technologies": {
                "machine learning": ["ml", "machine learning", "artificial intelligence", "ai", "deep learning"],
                "data science": ["data science", "data analysis", "analytics", "statistics"],
                "cloud": ["aws", "azure", "gcp", "google cloud", "cloud computing"],
                "docker": ["docker", "containerization", "kubernetes", "k8s"],
                "git": ["git", "github", "gitlab", "version control"]
            },
            "soft_skills": {
                "leadership": ["leadership", "team lead", "management", "mentoring"],
                "communication": ["communication", "presentation", "public speaking"],
                "problem solving": ["problem solving", "analytical", "critical thinking"]
            }
        }
        
        try:
            mappings_path = "config/skill_mappings.json"
            if os.path.exists(mappings_path):
                with open(mappings_path, 'r') as f:
                    return json.load(f)
            else:
                # Create default mappings file
                with open(mappings_path, 'w') as f:
                    json.dump(default_mappings, f, indent=2)
                return default_mappings
        except Exception as e:
            print_lg(f"Error loading skill mappings: {e}")
            return default_mappings
    
    def _load_keyword_weights(self) -> Dict:
        """Load keyword importance weights."""
        default_weights = {
            "required": 3.0,
            "preferred": 2.0,
            "nice_to_have": 1.0,
            "title_keywords": 2.5,
            "company_keywords": 1.5
        }
        
        try:
            weights_path = "config/keyword_weights.json"
            if os.path.exists(weights_path):
                with open(weights_path, 'r') as f:
                    return json.load(f)
            else:
                with open(weights_path, 'w') as f:
                    json.dump(default_weights, f, indent=2)
                return default_weights
        except Exception as e:
            print_lg(f"Error loading keyword weights: {e}")
            return default_weights
    
    def analyze_job_requirements(self, job_description: str, job_title: str, company: str) -> JobRequirements:
        """
        Analyzes job requirements from job description and title.
        """
        print_lg("🔍 Analyzing job requirements...")
        
        # Extract skills using pattern matching and AI
        required_skills = self._extract_required_skills(job_description)
        preferred_skills = self._extract_preferred_skills(job_description)
        
        # Determine experience level
        experience_level = self._extract_experience_level(job_description)
        
        # Extract education requirements
        education_requirements = self._extract_education_requirements(job_description)
        
        # Extract important keywords
        keywords = self._extract_keywords(job_description, job_title, company)
        
        # Determine industry and job function
        industry = self._determine_industry(company, job_description)
        job_function = self._determine_job_function(job_title, job_description)
        
        return JobRequirements(
            required_skills=required_skills,
            preferred_skills=preferred_skills,
            experience_level=experience_level,
            education_requirements=education_requirements,
            keywords=keywords,
            industry=industry,
            job_function=job_function
        )
    
    def _extract_required_skills(self, job_description: str) -> List[str]:
        """Extract required skills from job description."""
        skills = []
        
        # Look for required skills patterns
        required_patterns = [
            r"required:?\s*(.+?)(?:\n|\.|\;)",
            r"must have:?\s*(.+?)(?:\n|\.|\;)",
            r"essential:?\s*(.+?)(?:\n|\.|\;)",
            r"requirements:?\s*(.+?)(?:\n|\.|\;)"
        ]
        
        for pattern in required_patterns:
            matches = re.findall(pattern, job_description, re.IGNORECASE | re.DOTALL)
            for match in matches:
                skills.extend(self._parse_skills_from_text(match))
        
        return list(set(skills))
    
    def _extract_preferred_skills(self, job_description: str) -> List[str]:
        """Extract preferred skills from job description."""
        skills = []
        
        # Look for preferred skills patterns
        preferred_patterns = [
            r"preferred:?\s*(.+?)(?:\n|\.|\;)",
            r"nice to have:?\s*(.+?)(?:\n|\.|\;)",
            r"bonus:?\s*(.+?)(?:\n|\.|\;)",
            r"plus:?\s*(.+?)(?:\n|\.|\;)"
        ]
        
        for pattern in preferred_patterns:
            matches = re.findall(pattern, job_description, re.IGNORECASE | re.DOTALL)
            for match in matches:
                skills.extend(self._parse_skills_from_text(match))
        
        return list(set(skills))
    
    def _parse_skills_from_text(self, text: str) -> List[str]:
        """Parse individual skills from a text block."""
        # Split by common delimiters
        delimiters = [',', ';', '•', '-', '\n', 'and', '&']
        skills = [text]
        
        for delimiter in delimiters:
            new_skills = []
            for skill in skills:
                new_skills.extend([s.strip() for s in skill.split(delimiter)])
            skills = new_skills
        
        # Clean and filter skills
        cleaned_skills = []
        for skill in skills:
            skill = skill.strip().lower()
            if len(skill) > 2 and skill not in ['experience', 'knowledge', 'ability']:
                cleaned_skills.append(skill)
        
        return cleaned_skills
    
    def _extract_experience_level(self, job_description: str) -> str:
        """Extract experience level from job description."""
        experience_patterns = {
            'entry': r'(entry.level|junior|0.2\s+years?|new\s+grad)',
            'mid': r'(3.5\s+years?|mid.level|intermediate)',
            'senior': r'(5\+?\s+years?|senior|lead|principal)',
            'executive': r'(10\+?\s+years?|director|vp|executive|c.level)'
        }
        
        for level, pattern in experience_patterns.items():
            if re.search(pattern, job_description, re.IGNORECASE):
                return level
        
        return 'mid'  # Default to mid-level
    
    def _extract_education_requirements(self, job_description: str) -> List[str]:
        """Extract education requirements."""
        education = []
        
        education_patterns = [
            r"bachelor'?s?\s+degree",
            r"master'?s?\s+degree",
            r"phd|doctorate",
            r"associate'?s?\s+degree",
            r"high\s+school|diploma"
        ]
        
        for pattern in education_patterns:
            if re.search(pattern, job_description, re.IGNORECASE):
                education.append(pattern.replace(r'\s+', ' ').replace("'?s?", "'s"))
        
        return education
    
    def _extract_keywords(self, job_description: str, job_title: str, company: str) -> List[str]:
        """Extract important keywords for ATS optimization."""
        keywords = []
        
        # Add job title keywords
        keywords.extend(job_title.lower().split())
        
        # Add company name
        keywords.append(company.lower())
        
        # Extract technical terms and acronyms
        technical_pattern = r'\b[A-Z]{2,}\b|\b\w+\.\w+\b'
        technical_terms = re.findall(technical_pattern, job_description)
        keywords.extend([term.lower() for term in technical_terms])
        
        # Extract common job-related keywords
        job_keywords = [
            'agile', 'scrum', 'devops', 'ci/cd', 'api', 'rest', 'microservices',
            'database', 'sql', 'nosql', 'cloud', 'aws', 'azure', 'docker',
            'kubernetes', 'git', 'testing', 'automation'
        ]
        
        for keyword in job_keywords:
            if keyword in job_description.lower():
                keywords.append(keyword)
        
        return list(set(keywords))
    
    def _determine_industry(self, company: str, job_description: str) -> str:
        """Determine the industry based on company and job description."""
        industry_keywords = {
            'technology': ['software', 'tech', 'ai', 'machine learning', 'data'],
            'finance': ['bank', 'financial', 'investment', 'trading', 'fintech'],
            'healthcare': ['health', 'medical', 'hospital', 'pharmaceutical'],
            'retail': ['retail', 'e-commerce', 'consumer', 'shopping'],
            'consulting': ['consulting', 'advisory', 'strategy']
        }
        
        text = (company + ' ' + job_description).lower()
        
        for industry, keywords in industry_keywords.items():
            if any(keyword in text for keyword in keywords):
                return industry
        
        return 'general'
    
    def _determine_job_function(self, job_title: str, job_description: str) -> str:
        """Determine the job function."""
        function_keywords = {
            'engineering': ['engineer', 'developer', 'programmer', 'architect'],
            'data': ['data scientist', 'analyst', 'data engineer', 'ml engineer'],
            'product': ['product manager', 'product owner', 'pm'],
            'design': ['designer', 'ux', 'ui', 'design'],
            'marketing': ['marketing', 'growth', 'digital marketing'],
            'sales': ['sales', 'account manager', 'business development']
        }
        
        text = (job_title + ' ' + job_description).lower()
        
        for function, keywords in function_keywords.items():
            if any(keyword in text for keyword in keywords):
                return function
        
        return 'general'
    
    def optimize_resume_for_job(self, job_requirements: JobRequirements,
                               job_title: str, company: str) -> str:
        """
        Creates an AI-optimized resume for a specific job using Groq.
        """
        print_lg(f"🎯 AI-optimizing resume for {job_title} at {company}")

        # Parse base resume if not cached
        if not self.parsed_resume_cache:
            self.parsed_resume_cache = self._parse_base_resume()

        # Use AI optimization if available
        if self.groq_client:
            optimized_content = self._ai_optimize_resume_content(
                self.parsed_resume_cache, job_requirements, job_title, company
            )
        else:
            optimized_content = self.parsed_resume_cache

        # Create optimized resume document
        optimized_resume = self._create_optimized_resume(
            optimized_content, job_requirements, job_title, company
        )

        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        company_clean = re.sub(r'[^\w\s-]', '', company).strip()[:20]
        filename = f"resume_{company_clean}_{timestamp}.docx"
        filepath = os.path.join(self.optimized_resumes_dir, filename)

        # Save optimized resume
        optimized_resume.save(filepath)

        print_lg(f"✅ AI-optimized resume saved: {filepath}")
        return filepath

    def _ai_optimize_resume_content(self, resume_content: Dict, job_requirements: JobRequirements,
                                   job_title: str, company: str) -> Dict:
        """
        Use Groq AI to optimize resume content for specific job.
        """
        print_lg("🤖 Using Groq AI to optimize resume content...")

        try:
            # Prepare job description for AI
            job_description = f"""
            Job Title: {job_title}
            Company: {company}
            Required Skills: {', '.join(job_requirements.required_skills)}
            Preferred Skills: {', '.join(job_requirements.preferred_skills)}
            Experience Level: {job_requirements.experience_level}
            Industry: {job_requirements.industry}
            """

            # Use Groq AI to optimize resume content
            optimized_content = self.groq_client.optimize_resume_content(
                resume_content, job_description
            )

            print_lg("✅ Resume content optimized with AI")
            return optimized_content

        except Exception as e:
            print_lg(f"⚠️ AI optimization failed, using standard optimization: {e}")
            return resume_content

    def generate_ai_cover_letter(self, job_requirements: JobRequirements,
                                job_title: str, company: str, user_profile: Dict = None) -> str:
        """
        Generate AI-powered cover letter using Groq.
        """
        print_lg(f"📝 Generating AI cover letter for {job_title} at {company}")

        if not self.groq_client:
            return self._generate_fallback_cover_letter(job_title, company, user_profile)

        try:
            # Prepare job description
            job_description = f"""
            Job Title: {job_title}
            Company: {company}
            Required Skills: {', '.join(job_requirements.required_skills)}
            Preferred Skills: {', '.join(job_requirements.preferred_skills)}
            Experience Level: {job_requirements.experience_level}
            Industry: {job_requirements.industry}
            Job Function: {job_requirements.job_function}
            """

            # Prepare user profile
            if not user_profile:
                user_profile = {
                    'name': 'Job Applicant',
                    'experience_years': 3,
                    'skills': job_requirements.required_skills[:5],
                    'education': 'Bachelor\'s Degree',
                    'career_goals': [job_title]
                }

            # Generate cover letter with Groq AI
            cover_letter = self.groq_client.generate_cover_letter(
                job_description, user_profile, {'name': company}
            )

            print_lg("✅ AI cover letter generated successfully")
            return cover_letter

        except Exception as e:
            print_lg(f"⚠️ AI cover letter generation failed: {e}")
            return self._generate_fallback_cover_letter(job_title, company, user_profile)

    def _generate_fallback_cover_letter(self, job_title: str, company: str, user_profile: Dict = None) -> str:
        """Generate fallback cover letter if AI fails."""
        name = user_profile.get('name', 'Job Applicant') if user_profile else 'Job Applicant'

        return f"""Dear Hiring Manager,

I am writing to express my strong interest in the {job_title} position at {company}. With my background and experience, I am confident that I would be a valuable addition to your team.

My skills and experience align well with the requirements for this role, and I am particularly excited about the opportunity to contribute to {company}'s continued success. I am eager to bring my expertise and enthusiasm to your organization.

I would welcome the opportunity to discuss how my qualifications can benefit your team. Thank you for considering my application, and I look forward to hearing from you.

Best regards,
{name}"""
    
    def _parse_base_resume(self) -> Dict:
        """Parse the base resume into structured data."""
        print_lg("📄 Parsing base resume...")
        
        try:
            doc = Document(self.base_resume_path)
            
            sections = {
                'contact': [],
                'summary': [],
                'experience': [],
                'education': [],
                'skills': [],
                'projects': [],
                'certifications': []
            }
            
            current_section = 'contact'
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Detect section headers
                text_lower = text.lower()
                if any(header in text_lower for header in ['experience', 'work experience']):
                    current_section = 'experience'
                elif any(header in text_lower for header in ['education', 'academic']):
                    current_section = 'education'
                elif any(header in text_lower for header in ['skills', 'technical skills']):
                    current_section = 'skills'
                elif any(header in text_lower for header in ['projects', 'project']):
                    current_section = 'projects'
                elif any(header in text_lower for header in ['summary', 'objective']):
                    current_section = 'summary'
                elif any(header in text_lower for header in ['certification', 'certificate']):
                    current_section = 'certifications'
                else:
                    sections[current_section].append(text)
            
            return sections
            
        except Exception as e:
            print_lg(f"Error parsing base resume: {e}")
            return self._get_default_resume_structure()
    
    def _get_default_resume_structure(self) -> Dict:
        """Get default resume structure if parsing fails."""
        return {
            'contact': ['John Doe', 'john.doe@email.com', '+1234567890'],
            'summary': ['Experienced professional with strong technical skills'],
            'experience': ['Software Engineer at Tech Company (2020-Present)'],
            'education': ['Bachelor of Science in Computer Science'],
            'skills': ['Python', 'JavaScript', 'SQL', 'Git'],
            'projects': ['Personal Project: Web Application'],
            'certifications': []
        }
    
    def _create_optimized_resume(self, base_resume: Dict, job_requirements: JobRequirements,
                                job_title: str, company: str) -> Document:
        """Create an optimized resume document."""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Add contact information
        self._add_contact_section(doc, base_resume['contact'])
        
        # Add optimized summary
        self._add_optimized_summary(doc, base_resume['summary'], job_requirements, job_title)
        
        # Add optimized skills section
        self._add_optimized_skills(doc, base_resume['skills'], job_requirements)
        
        # Add experience section
        self._add_experience_section(doc, base_resume['experience'], job_requirements)
        
        # Add education section
        self._add_education_section(doc, base_resume['education'])
        
        # Add projects if relevant
        if base_resume['projects']:
            self._add_projects_section(doc, base_resume['projects'], job_requirements)
        
        # Add certifications if any
        if base_resume['certifications']:
            self._add_certifications_section(doc, base_resume['certifications'])
        
        return doc
    
    def _add_contact_section(self, doc: Document, contact_info: List[str]):
        """Add contact information section."""
        for info in contact_info[:3]:  # Limit to first 3 items
            p = doc.add_paragraph(info)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_optimized_summary(self, doc: Document, summary: List[str], 
                              job_requirements: JobRequirements, job_title: str):
        """Add optimized summary section."""
        doc.add_heading('Professional Summary', level=1)
        
        # Create optimized summary incorporating job keywords
        base_summary = ' '.join(summary) if summary else "Experienced professional"
        
        # Add job-specific keywords
        optimized_summary = f"{base_summary} with expertise in {', '.join(job_requirements.required_skills[:3])}. "
        optimized_summary += f"Seeking to leverage skills in {job_title.lower()} role."
        
        doc.add_paragraph(optimized_summary)
    
    def _add_optimized_skills(self, doc: Document, skills: List[str], 
                             job_requirements: JobRequirements):
        """Add optimized skills section."""
        doc.add_heading('Technical Skills', level=1)
        
        # Prioritize skills that match job requirements
        all_skills = skills.copy()
        
        # Add missing required skills if they're in our skill mappings
        for req_skill in job_requirements.required_skills:
            if req_skill not in [s.lower() for s in all_skills]:
                # Check if we have this skill under a different name
                for category, mappings in self.skill_mappings.items():
                    for skill, synonyms in mappings.items():
                        if req_skill in synonyms and skill not in [s.lower() for s in all_skills]:
                            all_skills.append(skill.title())
        
        # Sort skills by relevance to job
        sorted_skills = self._sort_skills_by_relevance(all_skills, job_requirements)
        
        # Format skills nicely
        skills_text = ' • '.join(sorted_skills)
        doc.add_paragraph(skills_text)
    
    def _sort_skills_by_relevance(self, skills: List[str], job_requirements: JobRequirements) -> List[str]:
        """Sort skills by relevance to job requirements."""
        skill_scores = {}
        
        for skill in skills:
            score = 0
            skill_lower = skill.lower()
            
            # Check if skill is required
            if skill_lower in [s.lower() for s in job_requirements.required_skills]:
                score += 10
            
            # Check if skill is preferred
            if skill_lower in [s.lower() for s in job_requirements.preferred_skills]:
                score += 5
            
            # Check if skill appears in keywords
            if skill_lower in [k.lower() for k in job_requirements.keywords]:
                score += 3
            
            skill_scores[skill] = score
        
        # Sort by score (descending) and return
        return sorted(skills, key=lambda x: skill_scores.get(x, 0), reverse=True)
    
    def _add_experience_section(self, doc: Document, experience: List[str], 
                               job_requirements: JobRequirements):
        """Add experience section with job-relevant emphasis."""
        doc.add_heading('Professional Experience', level=1)
        
        for exp in experience:
            doc.add_paragraph(exp, style='List Bullet')
    
    def _add_education_section(self, doc: Document, education: List[str]):
        """Add education section."""
        doc.add_heading('Education', level=1)
        
        for edu in education:
            doc.add_paragraph(edu)
    
    def _add_projects_section(self, doc: Document, projects: List[str], 
                             job_requirements: JobRequirements):
        """Add projects section."""
        doc.add_heading('Projects', level=1)
        
        for project in projects:
            doc.add_paragraph(project, style='List Bullet')
    
    def _add_certifications_section(self, doc: Document, certifications: List[str]):
        """Add certifications section."""
        doc.add_heading('Certifications', level=1)
        
        for cert in certifications:
            doc.add_paragraph(cert, style='List Bullet')
    
    def calculate_resume_match_score(self, resume_path: str, job_requirements: JobRequirements) -> float:
        """Calculate how well a resume matches job requirements."""
        try:
            doc = Document(resume_path)
            resume_text = ' '.join([p.text for p in doc.paragraphs]).lower()
            
            score = 0.0
            total_possible = 0
            
            # Check required skills
            for skill in job_requirements.required_skills:
                total_possible += self.keyword_weights['required']
                if skill.lower() in resume_text:
                    score += self.keyword_weights['required']
            
            # Check preferred skills
            for skill in job_requirements.preferred_skills:
                total_possible += self.keyword_weights['preferred']
                if skill.lower() in resume_text:
                    score += self.keyword_weights['preferred']
            
            # Check keywords
            for keyword in job_requirements.keywords:
                total_possible += self.keyword_weights['nice_to_have']
                if keyword.lower() in resume_text:
                    score += self.keyword_weights['nice_to_have']
            
            return (score / total_possible) if total_possible > 0 else 0.0
            
        except Exception as e:
            print_lg(f"Error calculating resume match score: {e}")
            return 0.0
