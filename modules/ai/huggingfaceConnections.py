# Author: Manideep Reddy Eevuri
# Hugging Face AI Integration

from config.secrets import *
from config.settings import showAiErrorAlerts
from modules.helpers import print_lg, critical_error_log, convert_to_json
from modules.ai.prompts import *

import requests
from typing import Literal

apiCheckInstructions = """
1. Make sure your Hugging Face API key is correct.
2. Check your model name and endpoint.
3. Open `secrets.py` in `/config` folder to configure your Hugging Face API connection.
ERROR:
"""

def huggingface_create_client():
    """
    Returns a dummy client (just the API key) for compatibility.
    """
    if not use_AI:
        raise ValueError("AI is not enabled! Please enable it by setting `use_AI = True` in `secrets.py` in `config` folder.")
    if not huggingface_api_key or huggingface_api_key in ["", "not-needed"]:
        raise ValueError("Hugging Face API key is not set!")
    print_lg("---- SUCCESSFULLY INITIALIZED HUGGING FACE CLIENT! ----")
    print_lg(f"Using Hugging Face API Key: {huggingface_api_key[:8]}... (hidden)")
    return huggingface_api_key

def huggingface_inference(prompt: str, model: str = "gpt2", temperature: float = 0.7, max_tokens: int = 512) -> str:
    """
    Calls the Hugging Face Inference API for text generation.
    """
    url = f"https://api-inference.huggingface.co/models/{model}"
    headers = {"Authorization": f"Bearer {huggingface_api_key}"}
    payload = {"inputs": prompt, "parameters": {"temperature": temperature, "max_new_tokens": max_tokens}}
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=60)
        response.raise_for_status()
        result = response.json()
        if isinstance(result, list) and len(result) > 0 and "generated_text" in result[0]:
            return result[0]["generated_text"]
        elif isinstance(result, dict) and "error" in result:
            raise ValueError(result["error"])
        else:
            return str(result)
    except Exception as e:
        critical_error_log("Error occurred while calling Hugging Face Inference API.", e)
        return f"Error: {e}"

def huggingface_answer_question(client, question: str, options: list[str] | None = None, question_type: Literal['text', 'textarea', 'single_select', 'multiple_select'] = 'text', job_description: str = None, about_company: str = None, user_information_all: str = None, model: str = "gpt2") -> str:
    prompt = ai_answer_prompt.format(user_information_all or "N/A", question)
    if job_description:
        prompt += f"\nJob Description:\n{job_description}"
    if about_company:
        prompt += f"\nAbout the Company:\n{about_company}"
    return huggingface_inference(prompt, model=model)

def huggingface_generate_resume(client, job_description: str, about_company: str, required_skills: dict, user_information: str = None, model: str = "gpt2", save_path: str = None) -> str:
    from docx import Document
    from datetime import datetime
    import os
    from modules.ai.prompts import ai_resume_prompt
    prompt = ai_resume_prompt.format(
        user_information=user_information or "N/A",
        job_description=job_description or "N/A",
        about_company=about_company or "N/A",
        required_skills=str(required_skills) if required_skills else "N/A"
    )
    response = huggingface_inference(prompt, model=model)
    if isinstance(response, ValueError):
        return response
    doc = Document()
    lines = response.split('\n')
    for line in lines:
        doc.add_paragraph(line)
    if not save_path:
        save_path = f"all resumes/Customized_Resume_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    doc.save(save_path)
    print_lg(f"Customized resume saved to: {save_path}")
    return save_path

def huggingface_generate_coverletter(client, job_description: str, about_company: str, required_skills: dict, user_information: str = None, model: str = "gpt2", save_path: str = None) -> str:
    from docx import Document
    from datetime import datetime
    import os
    from modules.ai.prompts import ai_coverletter_prompt
    prompt = ai_coverletter_prompt.format(
        user_information=user_information or "N/A",
        job_description=job_description or "N/A",
        about_company=about_company or "N/A",
        required_skills=str(required_skills) if required_skills else "N/A"
    )
    response = huggingface_inference(prompt, model=model)
    if isinstance(response, ValueError):
        return response
    doc = Document()
    lines = response.split('\n')
    for line in lines:
        doc.add_paragraph(line)
    if not save_path:
        save_path = f"all resumes/Customized_CoverLetter_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    doc.save(save_path)
    print_lg(f"Cover letter saved to: {save_path}")
    return save_path 