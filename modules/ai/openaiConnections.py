# Author: Manideep Reddy Eevuri
# LinkedIn: https://www.linkedin.com/in/manideep-reddy-eevuri-661659268/
# GitHub: https://github.com/Maniredii
#

from config.secrets import *
from config.settings import showAiErrorAlerts
from config.personals import ethnicity, gender, disability_status, veteran_status
from config.questions import *
from config.search import security_clearance, did_masters

from modules.helpers import print_lg, critical_error_log, convert_to_json
from modules.ai.prompts import *

from pyautogui import confirm
from openai import OpenAI
from openai.types.model import Model
from openai.types.chat import ChatCompletion, ChatCompletionChunk
from typing import Iterator, Literal


apiCheckInstructions = """

1. Make sure your AI API connection details like url, key, model names, etc are correct.
2. If you're using an local LLM, please check if the server is running.
3. Check if appropriate LLM and Embedding models are loaded and running.

Open `secret.py` in `/config` folder to configure your AI API connections.

ERROR:
"""

# Function to show an AI error alert
def ai_error_alert(message: str, stackTrace: str, title: str = "AI Connection Error") -> None:
    """
    Function to show an AI error alert and log it.
    """
    global showAiErrorAlerts
    if showAiErrorAlerts:
        if "Pause AI error alerts" == confirm(f"{message}{stackTrace}\n", title, ["Pause AI error alerts", "Okay Continue"]):
            showAiErrorAlerts = False
    critical_error_log(message, stackTrace)


# Function to check if an error occurred
def ai_check_error(response: ChatCompletion | ChatCompletionChunk) -> None:
    """
    Function to check if an error occurred.
    * Takes in `response` of type `ChatCompletion` or `ChatCompletionChunk`
    * Raises a `ValueError` if an error is found
    """
    if response.model_extra.get("error"):
        raise ValueError(
            f'Error occurred with API: "{response.model_extra.get("error")}"'
        )


# Function to create an OpenAI client
def ai_create_openai_client() -> OpenAI:
    """
    Function to create an OpenAI client.
    * Takes no arguments
    * Returns an `OpenAI` object
    """
    try:
        print_lg("Creating OpenAI client...")
        if not use_AI:
            raise ValueError("AI is not enabled! Please enable it by setting `use_AI = True` in `secrets.py` in `config` folder.")
        
        client = OpenAI(base_url=llm_api_url, api_key=llm_api_key)

        models = ai_get_models_list(client)
        if "error" in models:
            raise ValueError(models[1])
        if len(models) == 0:
            raise ValueError("No models are available!")
        if llm_model not in [model.id for model in models]:
            raise ValueError(f"Model `{llm_model}` is not found!")
        
        print_lg("---- SUCCESSFULLY CREATED OPENAI CLIENT! ----")
        print_lg(f"Using API URL: {llm_api_url}")
        print_lg(f"Using Model: {llm_model}")
        print_lg("Check './config/secrets.py' for more details.\n")
        print_lg("---------------------------------------------")

        return client
    except Exception as e:
        ai_error_alert(f"Error occurred while creating OpenAI client. {apiCheckInstructions}", e)


# Function to close an OpenAI client
def ai_close_openai_client(client: OpenAI) -> None:
    """
    Function to close an OpenAI client.
    * Takes in `client` of type `OpenAI`
    * Returns no value
    """
    try:
        if client:
            print_lg("Closing OpenAI client...")
            client.close()
    except Exception as e:
        ai_error_alert("Error occurred while closing OpenAI client.", e)



# Function to get list of models available in OpenAI API
def ai_get_models_list(client: OpenAI) -> list[ Model | str]:
    """
    Function to get list of models available in OpenAI API.
    * Takes in `client` of type `OpenAI`
    * Returns a `list` object
    """
    try:
        print_lg("Getting AI models list...")
        if not client: raise ValueError("Client is not available!")
        models = client.models.list()
        ai_check_error(models)
        print_lg("Available models:")
        print_lg(models.data, pretty=True)
        return models.data
    except Exception as e:
        critical_error_log("Error occurred while getting models list!", e)
        return ["error", e]

def model_supports_temperature(model_name: str) -> bool:
    """
    Checks if the specified model supports the temperature parameter.
    
    Args:
        model_name (str): The name of the AI model.
    
    Returns:
        bool: True if the model supports temperature adjustments, otherwise False.
    """
    return model_name in ["gpt-3.5-turbo", "gpt-4", "gpt-4-turbo", "gpt-4o", "gpt-4o-mini"]

# Function to get chat completion from OpenAI API
def ai_completion(client: OpenAI, messages: list[dict], response_format: dict = None, temperature: float = 0, stream: bool = stream_output) -> dict | ValueError:
    """
    Function that completes a chat and prints and formats the results of the OpenAI API calls.
    * Takes in `client` of type `OpenAI`
    * Takes in `messages` of type `list[dict]`. Example: `[{"role": "user", "content": "Hello"}]`
    * Takes in `response_format` of type `dict` for JSON representation, default is `None`
    * Takes in `temperature` of type `float` for temperature, default is `0`
    * Takes in `stream` of type `bool` to indicate if it's a streaming call or not
    * Returns a `dict` object representing JSON response, will try to convert to JSON if `response_format` is given
    """
    if not client: raise ValueError("Client is not available!")

    params = {"model": llm_model, "messages": messages, "stream": stream}

    if model_supports_temperature(llm_model):
        params["temperature"] = temperature
    if response_format and llm_spec in ["openai", "openai-like"]:
        params["response_format"] = response_format

    completion = client.chat.completions.create(**params)

    result = ""
    
    # Log response
    if stream:
        print_lg("--STREAMING STARTED")
        for chunk in completion:
            ai_check_error(chunk)
            chunkMessage = chunk.choices[0].delta.content
            if chunkMessage != None:
                result += chunkMessage
            print_lg(chunkMessage, end="", flush=True)
        print_lg("\n--STREAMING COMPLETE")
    else:
        ai_check_error(completion)
        result = completion.choices[0].message.content
    
    if response_format:
        result = convert_to_json(result)
    
    print_lg("\nAI Answer to Question:\n")
    print_lg(result, pretty=response_format)
    return result


def ai_extract_skills(client: OpenAI, job_description: str, stream: bool = stream_output) -> dict | ValueError:
    """
    Function to extract skills from job description using OpenAI API.
    * Takes in `client` of type `OpenAI`
    * Takes in `job_description` of type `str`
    * Takes in `stream` of type `bool` to indicate if it's a streaming call
    * Returns a `dict` object representing JSON response
    """
    print_lg("-- EXTRACTING SKILLS FROM JOB DESCRIPTION")
    try:        
        prompt = extract_skills_prompt.format(job_description)

        messages = [{"role": "user", "content": prompt}]
        ##> ------ Dheeraj Deshwal : dheeraj20194@iiitd.ac.in/dheerajdeshwal9811@gmail.com - Bug fix ------
        return ai_completion(client, messages, response_format=extract_skills_response_format, stream=stream)
    ##<
    except Exception as e:
        ai_error_alert(f"Error occurred while extracting skills from job description. {apiCheckInstructions}", e)


##> ------ Dheeraj Deshwal : dheeraj9811 Email:dheeraj20194@iiitd.ac.in/dheerajdeshwal9811@gmail.com - Feature ------
def ai_answer_question(
    client: OpenAI, 
    question: str, options: list[str] | None = None, question_type: Literal['text', 'textarea', 'single_select', 'multiple_select'] = 'text', 
    job_description: str = None, about_company: str = None, user_information_all: str = None,
    stream: bool = stream_output
) -> dict | ValueError:
    """
    Function to generate AI-based answers for questions in a form.
    
    Parameters:
    - `client`: OpenAI client instance.
    - `question`: The question being answered.
    - `options`: List of options (for `single_select` or `multiple_select` questions).
    - `question_type`: Type of question (text, textarea, single_select, multiple_select) It is restricted to one of four possible values.
    - `job_description`: Optional job description for context.
    - `about_company`: Optional company details for context.
    - `user_information_all`: information about you, AI cna use to answer question eg: Resume-like user information.
    - `stream`: Whether to use streaming AI completion.
    
    Returns:
    - `str`: The AI-generated answer.
    """

    print_lg("-- ANSWERING QUESTION using AI")
    try:
        prompt = ai_answer_prompt.format(user_information_all or "N/A", question)
         # Append optional details if provided
        if job_description and job_description != "Unknown":
            prompt += f"\nJob Description:\n{job_description}"
        if about_company and about_company != "Unknown":
            prompt += f"\nAbout the Company:\n{about_company}"

        messages = [{"role": "user", "content": prompt}]
        print_lg("Prompt we are passing to AI: ", prompt)
        response =  ai_completion(client, messages, stream=stream)
        # print_lg("Response from AI: ", response)
        return response
    except Exception as e:
        ai_error_alert(f"Error occurred while answering question. {apiCheckInstructions}", e)
##<


def ai_gen_experience(
    client: OpenAI, 
    job_description: str, about_company: str, 
    required_skills: dict, user_experience: dict,
    stream: bool = stream_output
) -> dict | ValueError:
    pass



def ai_generate_resume(
    client: OpenAI, 
    job_description: str, about_company: str, required_skills: dict, user_information: str = None,
    stream: bool = stream_output,
    save_path: str = None
) -> str | ValueError:
    '''
    Function to generate a customized resume using OpenAI based on job description and user information.
    '''
    from docx import Document
    from datetime import datetime
    import os
    
    try:
        from modules.ai.prompts import ai_resume_prompt
        
        # Create prompt for resume customization
        prompt = ai_resume_prompt.format(
            user_information=user_information or "N/A",
            job_description=job_description or "N/A",
            about_company=about_company or "N/A",
            required_skills=str(required_skills) if required_skills else "N/A"
        )
        
        messages = [{"role": "user", "content": prompt}]
        
        # Generate customized resume content
        response = ai_completion(client, messages, stream=stream)
        
        if isinstance(response, ValueError):
            return response
        
        # Create a new document with the customized content
        doc = Document()
        
        # Parse the AI response and create structured resume
        lines = response.split('\n')
        current_section = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Detect section headers
            if line.upper() in ['EDUCATION', 'EXPERIENCE', 'SKILLS', 'PROJECTS', 'CERTIFICATES', 'SUMMARY']:
                current_section = line
                doc.add_heading(line, level=1)
            elif line.startswith('##'):
                # Sub-section
                doc.add_heading(line.replace('#', '').strip(), level=2)
            elif line.startswith('#'):
                # Main section
                doc.add_heading(line.replace('#', '').strip(), level=1)
            else:
                # Regular content
                doc.add_paragraph(line)
        
        # Generate filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        company_name = about_company.split()[0] if about_company else "Company"
        filename = f"customized_resume_{company_name}_{timestamp}.docx"
        
        # Save to specified path or default location
        if save_path:
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            file_path = os.path.join(save_path, filename)
        else:
            # Create generated_resumes directory if it doesn't exist
            os.makedirs("generated_resumes", exist_ok=True)
            file_path = f"generated_resumes/{filename}"
        
        doc.save(file_path)
        print_lg(f"Customized resume saved: {file_path}")
        
        return file_path
        
    except Exception as e:
        print_lg(f"Error generating customized resume: {e}")
        return ValueError(f"Failed to generate resume: {e}")


def ai_generate_coverletter(
    client: OpenAI, 
    job_description: str, about_company: str, required_skills: dict, user_information: str = None,
    stream: bool = stream_output,
    save_path: str = None
) -> str | ValueError:
    '''
    Function to generate a cover letter using OpenAI and save it as a .docx file.
    '''
    from docx import Document
    from datetime import datetime
    import os
    
    try:
        from modules.ai.prompts import ai_cover_letter_prompt
        
        prompt = ai_cover_letter_prompt.format(
            user_information=user_information or "N/A",
            job_description=job_description or "N/A",
            about_company=about_company or "N/A"
        )
        
        messages = [{"role": "user", "content": prompt}]
        cover_letter = ai_completion(client, messages, stream=stream)
        
        if isinstance(cover_letter, ValueError):
            return cover_letter
        
        # Save to .docx
        doc = Document()
        doc.add_heading('Cover Letter', 0)
        for para in cover_letter.split('\n'):
            if para.strip():
                doc.add_paragraph(para.strip())
        
        if not save_path:
            now = datetime.now().strftime('%Y%m%d_%H%M%S')
            save_path = f"generated_cover_letters/cover_letter_{now}.docx"
        
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        doc.save(save_path)
        print_lg(f"Cover letter saved to {save_path}")
        return save_path
        
    except Exception as e:
        ai_error_alert(f"Error occurred while generating cover letter. {apiCheckInstructions}", e)
        return ValueError(str(e))



##< Evaluation Agents
def ai_evaluate_resume(
    client: OpenAI, 
    job_description: str, about_company: str, required_skills: dict,
    resume: str,
    stream: bool = stream_output
) -> dict | ValueError:
    pass



def ai_evaluate_resume(
    client: OpenAI, 
    job_description: str, about_company: str, required_skills: dict,
    resume: str,
    stream: bool = stream_output
) -> dict | ValueError:
    pass



def ai_check_job_relevance(
    client: OpenAI, 
    job_description: str, about_company: str,
    stream: bool = stream_output
) -> dict:
    pass
#>